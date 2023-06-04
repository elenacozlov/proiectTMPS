from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
import docx
from docx.shared import Pt

class DocumentGenerator:
    def __init__(self, strategy):
        self.strategy = strategy

    def generate(self, lines, entry_values):
        self.strategy.generate(lines, entry_values)

class DocumentGeneratorFacade:
    def __init__(self, output_dir):
        self.output_dir = output_dir

    def generate_pdf(self, lines, entries):
        strategy = PDFDocumentGeneratorStrategy(self.output_dir)
        DocumentGenerator(strategy).generate(lines, entries)

    def generate_word(self, lines, entries):
        strategy = WordDocumentGeneratorStrategy(self.output_dir)
        DocumentGenerator(strategy).generate(lines, entries)


class PDFDocumentGeneratorStrategy:
    def __init__(self, output_dir):
        self.output_dir = output_dir

    def generate(self, lines, entry_values):
        file_path = self.output_dir + "/cerere.pdf"

        doc = SimpleDocTemplate(file_path, pagesize=A4)
        story = []

        for line, alignment in lines:
            p = Paragraph(line, ParagraphStyle(name="Normal", alignment=alignment.reportlab()))
            story.append(p)
            story.append(Spacer(1, 12))

        doc.build(story)


class WordDocumentGeneratorStrategy:
    def __init__(self, output_dir):
        self.output_dir = output_dir

    def generate(self, lines, entry_values):
        file_path = self.output_dir + "/cerere.docx"

        doc = docx.Document()

        for text, alignment in lines:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            paragraph.alignment = alignment.docx()
            run.bold = True
            run.font.size = Pt(12)

        doc.save(file_path)